"""CV parsing: extract raw text from PDF/DOCX uploads."""
from __future__ import annotations

import io
from typing import BinaryIO

from pypdf import PdfReader
from docx import Document


def parse_pdf(file: BinaryIO) -> str:
    reader = PdfReader(file)
    chunks = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:
            continue
    return _clean("\n".join(chunks))


def parse_docx(file: BinaryIO) -> str:
    doc = Document(file)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells))
    return _clean("\n".join(paragraphs))


def parse_cv(uploaded_file) -> str:
    """Dispatch on filename extension. Accepts Streamlit UploadedFile."""
    name = (uploaded_file.name or "").lower()
    data = uploaded_file.read()
    buf = io.BytesIO(data)
    if name.endswith(".pdf"):
        return parse_pdf(buf)
    if name.endswith(".docx"):
        return parse_docx(buf)
    if name.endswith(".txt"):
        return _clean(data.decode("utf-8", errors="ignore"))
    raise ValueError(f"Unsupported file type: {name}. Use PDF, DOCX, or TXT.")


def _clean(text: str) -> str:
    lines = [ln.rstrip() for ln in text.splitlines()]
    # collapse excessive blank lines
    out, blank = [], 0
    for ln in lines:
        if not ln.strip():
            blank += 1
            if blank <= 1:
                out.append("")
        else:
            blank = 0
            out.append(ln)
    return "\n".join(out).strip()
